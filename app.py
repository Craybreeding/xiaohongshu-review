import streamlit as st
import re
import os
import json
from datetime import datetime
from docx import Document
import io
import urllib.request

TODAY = datetime.now().strftime("%Y%m%d")

# ========== 审核规则常量 ==========
REQUIRED_TAGS = [
    "#能恩全护", "#能恩全护水奶", "#适度水解", "#适度水解奶粉",
    "#适度水解奶粉推荐", "#防敏奶粉", "#第一口奶粉", "#雀巢适度水解"
]

TITLE_KEYWORDS = ["适度水解", "防敏", "科普"]
BODY_KEYWORDS = ["适度水解", "防敏", "能恩全护"]
COVER_KEYWORDS = ["适度水解", "防敏", "科普"]

FORBIDDEN_WORDS = {
    "禁止词": ["敏宝", "奶瓶", "奶嘴", "新生儿", "过敏", "疾病"],
    "禁疗效表述": ["预防", "生长", "发育", "免疫"],
    "禁绝对化": ["最", "第一", "TOP1"],
}

FORBIDDEN_EXCEPTIONS = {
    "第一": ["第一口奶粉", "第一口配方粉", "第一口奶", "第一口配方", "第一步"],
    "最": ["最近", "最后", "最终", "最初", "最多", "最大"],
    "过敏": ["过敏源", "过敏原", "过敏史", "过敏体质", "过敏风险", "过敏率", "谈敏色变", "怕娃过敏"],
    "预防": ["预防敏感"],
    "生长": ["生长指标", "生长曲线"],
    "新生儿": ["#新生儿奶粉", "#新生儿"],
}

FORBIDDEN_REPLACEMENTS = {
    "过敏": "敏敏", "敏宝": "敏感体质宝宝",
    "新生儿": "初生宝宝", "预防": "防敏",
    "生长": "成长", "发育": "噌噌长",
    "免疫": "保护力", "疾病": "不适",
}

# 必提需润色卖点 (4大方向10小方向)
PARAPHRASE_SELLING_POINTS = [
    {"category": "敏敏背景", "idx": 1,
     "text": "我国初生宝宝敏敏率高达40%，要是有父母敏敏史，宝宝敏敏的概率将飙升到80%",
     "fragment": "敏敏率高达40%"},
    {"category": "防敏-水解技术", "idx": 2,
     "text": "易敏的大分子牛奶蛋白切割成温和的适度水解小分子牛奶蛋白，精准去掉致敏片段的同时，又完整保留了蛋白有益营养",
     "fragment": "切割成温和的适度水解小分子"},
    {"category": "防敏-水解技术", "idx": 3,
     "text": "全球专业人士优先推荐呢",
     "fragment": "全球专业人士优先推荐"},
    {"category": "自护力", "idx": 4,
     "text": "6种HMO加上明星双菌B.Infantis 和 Bb-12，两者强强联合，协同作用释放高倍的原生保护力",
     "fragment": "两者强强联合"},
    {"category": "自护力", "idx": 5,
     "text": "短短28天就能调理好娃的肚肚菌菌环境，从肚肚到全身都建起坚固的防护屏障",
     "fragment": "从肚肚到全身"},
    {"category": "自护力", "idx": 6,
     "text": "保护力能持续15个月，助力娃成长",
     "fragment": "助力娃成长"},
    {"category": "自护力", "idx": 7,
     "text": "四维成长曲线特别出色",
     "fragment": "四维成长曲线"},
    {"category": "基础营养", "idx": 8,
     "text": "基础营养也很抗打",
     "fragment": "基础营养也很抗打"},
    {"category": "基础营养", "idx": 9,
     "text": "25种维生素和矿物质拉满",
     "fragment": "维生素和矿物质拉满"},
    {"category": "基础营养", "idx": 10,
     "text": "全乳糖的配方口味清淡，宝宝爱喝",
     "fragment": "全乳糖的配方口味清淡，宝宝爱喝"},
]

# 必提不可修改卖点 (3大切角10小切角)
FIXED_SELLING_POINTS = [
    {"category": "防敏-水解技术", "idx": 1, "text": "多项科学实证的雀巢尖峰水解技术"},
    {"category": "防敏-水解技术", "idx": 2, "text": "温和的适度水解小分子牛奶蛋白"},
    {"category": "防敏-水解技术", "idx": 3, "text": "防敏领域权威德国GINI研究认证，能长效防敏20年，相比于牛奶蛋白致敏性降低1000倍"},
    {"category": "自护力", "idx": 4, "text": "采用了全球创新的超倍自护科技"},
    {"category": "自护力", "idx": 5, "text": "6种HMO加上明星双菌B.Infantis 和 Bb-12"},
    {"category": "自护力", "idx": 6, "text": "协同作用释放高倍的原生保护力"},
    {"category": "自护力", "idx": 7, "text": "短短28天就能调理好娃的肚肚菌菌环境"},
    {"category": "自护力", "idx": 8, "text": "保护力能持续15个月"},
    {"category": "基础营养", "idx": 9, "text": "25种维生素和矿物质"},
    {"category": "基础营养", "idx": 10, "text": "全乳糖的配方口味清淡"},
]

# 允许删减的卖点
OPTIONAL_SELLING_POINTS = [
    {"category": "防敏-水解技术",
     "text": "欧盟认可及全球30+科学实证背书，硬实力真材实料摆出来!",
     "fragment": "欧盟认可"},
    {"category": "基础营养",
     "text": "早期配方还搭配了牛磺酸、胆碱、核苷酸等关键营养。不添加蔗、香精这些不友好成分。",
     "fragment": "牛磺酸、胆碱、核苷酸"},
]

# 卖点顺序锚点
ORDER_ANCHORS = {
    "防敏-水解技术": ["水解技术", "尖峰水解", "GINI", "致敏性降低", "适度水解小分子"],
    "自护力": ["超倍自护", "HMO", "双菌", "B.Infantis", "原生保护力"],
    "基础营养": ["维生素和矿物质", "全乳糖"],
}

# 标准卖点示例
SELLING_POINT_EXAMPLE = """我国初生宝宝敏敏率高达40%，要是有父母敏敏史，宝宝敏敏的概率将飙升到80%；敏敏高发的原因（未经产道挤压、养宠专业人士建议：不少专业人士建议，可以给宝宝选择适度水解配方粉作为宝宝的第一口配方粉

拥有多项科学实证的雀巢尖峰水解科技，就像给蛋白装了精准切割器，把易敏的大分子牛奶蛋白切割成温和的100%适度水解小分子牛奶蛋白，精准去掉致敏片段的同时，又完整保留了蛋白有益营养，更亲和宝宝娇肚肚!不仅有防敏领域权威德国GINI研究认证，能长效防敏20年，还有欧盟认可及全球30+科学实证背书，相比于牛奶蛋白致敏性降低1000倍，硬实力真材实料摆出来!怪不得全球专业人士优先推荐呢！/全球专业人士优先推荐是有道理的

它采用了全球创新的"超倍自护科技"，其中6种HMO加上明星双菌B.Infantis 和 Bb-12，两者强强联合，协同作用释放高倍的原生保护力！短短28天就能调理好娃的肚肚菌菌环境，从肚肚到全身都建起坚固的防护屏障。更关键的是，这份保护力能持续15个月，完美覆盖宝宝的黄金发育期，助力娃噌长、稳稳长~ 有它助力，娃的四维成长曲线特别出色!

基础营养也很抗打，25种维生素和矿物质拉满，早期配方还搭配了牛磺酸、胆碱、核苷酸等关键营养。全乳糖的配方口味清淡，不添加蔗、香精这些不友好成分，宝宝爱喝，妈妈放心。"""

# 人话修改 Prompt
RENHUA_PROMPT = """你是小红书顶级爆文写手，擅长把硬广写成真实分享。现在帮我改写【能恩全护奶粉】的KOL稿件。

⚠️ 【最重要的3个硬性要求 - 必须全部满足】⚠️
1. 正文字数必须在800-900字之间（这是最重要的！太短或太长都不行）
2. 必须有强烈的小红书活人感、爆文感、真实分享感
3. 必须包含下面10句话术（可以自然融入，但字字不能改）

【10句必须原封不动出现的话术】
① 多项科学实证的雀巢尖峰水解技术
② 温和的适度水解小分子牛奶蛋白
③ 防敏领域权威德国GINI研究认证，能长效防敏20年，相比于牛奶蛋白致敏性降低1000倍
④ 采用了全球创新的超倍自护科技
⑤ 6种HMO加上明星双菌B.Infantis 和 Bb-12
⑥ 协同作用释放高倍的原生保护力
⑦ 短短28天就能调理好娃的肚肚菌菌环境
⑧ 保护力能持续15个月
⑨ 25种维生素和矿物质
⑩ 全乳糖的配方口味清淡

【小红书爆文写法 - 这才是活人感！】
🔥 开头要炸：用"姐妹们！""救命！""后悔没早知道！"等情绪钩子开场
🔥 说人话：把"因此建议"换成"所以我真心推荐"，把"具有"换成"有"
🔥 像聊天：多用"我""你""咱家娃"，写得像在跟闺蜜分享经验
🔥 有情绪：加入"说实话""真的绝了""一开始我也担心"等真实感受
🔥 短句+emoji：每句话不超过20字，适当加💡✨🔥❗等emoji
🔥 有节奏：用"！"比"。"多，读起来要有激动感
🔥 结尾要互动："姐妹们冲！""有同款宝宝的妈妈评论区举手🙋‍♀️"

【内容结构】（按这个顺序写，自然过渡）
1. 开篇钩子：作为育婴师/营养师，说说妈妈们最担心的户外带娃敏敏问题（约70字）
2. 痛点共鸣：我国初生宝宝敏敏率40%，有家族史飙到80%，太可怕了（约70字）
3. 科学支招：第一口奶粉选对很关键，推荐适度水解配方（约200字）
4. 产品种草：重点介绍能恩全护的水解技术、自护力配方、营养成分（约450字，融入10句话术）
5. 收尾号召：想带娃放心玩，选对奶粉是第一步！（约60字）

【其他要求】
- 提供3个标题备选（必须含：适度水解、防敏、科普）
- 提供10个以上话题标签（必须含：#能恩全护 #适度水解 #适度水解奶粉推荐 #第一口奶粉）
- 卖点顺序：防敏水解技术 → 自护力 → 基础营养
- 禁词替换：过敏→敏敏，预防→防敏，新生儿→初生宝宝，生长发育→成长
- 绝对禁止出现：敏宝、奶瓶、奶嘴、疾病、治疗、免疫

【输出格式】
### 标题备选（3个）
1. xxx
2. xxx
3. xxx

### 正文（800-900字，必须写够！）
（这里输出完整正文，要有小红书爆文的活人感！）

### 话题标签
#能恩全护 #适度水解 ...（10个以上）

---
【需要改写的KOL原稿】
{content}"""

# ========== 工具函数 ==========
def read_docx(file):
    doc = Document(io.BytesIO(file.read()))
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
    return "\n".join(text)

def call_llm_api(prompt):
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return "Error: 未设置OPENAI_API_KEY环境变量。请在Render环境变量中设置。"
    url = "https://api.openai.com/v1/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
    data = {
        "model": "gpt-4o",
        "max_tokens": 4000,
        "temperature": 0.7,
        "messages": [
            {"role": "system", "content": "你是一个专业的小红书KOL稿件改写助手。严格遵守字数要求（800-900中文字）和话术要求。"},
            {"role": "user", "content": prompt}
        ]
    }
    req = urllib.request.Request(url, data=json.dumps(data).encode('utf-8'), headers=headers)
    try:
        with urllib.request.urlopen(req, timeout=180) as response:
            result = json.loads(response.read().decode('utf-8'))
            return result["choices"][0]["message"]["content"]
    except urllib.error.HTTPError as e:
        error_body = e.read().decode('utf-8') if e.fp else ""
        return f"Error: HTTP {e.code} - {error_body[:200]}"
    except urllib.error.URLError as e:
        return f"Error: 网络连接失败 - {str(e.reason)}"
    except Exception as e:
        return f"Error: {str(e)}"

def count_chinese(text):
    return len(re.findall(r'[\u4e00-\u9fff]', text))

def extract_tags(content):
    return re.findall(r'#[\w\u4e00-\u9fff]+', content)

def extract_title(content):
    for line in content.split('\n'):
        line = line.strip()
        if line and not line.startswith('#'):
            return line
    return ""

def detect_titles(content):
    """智能检测标题数量，支持多种格式"""
    titles = []

    # 格式1: ### 标题备选 后面的编号列表
    title_section = re.search(r'###\s*标题备选.*?\n(.*?)(?=###|$)', content, re.DOTALL)
    if title_section:
        numbered = re.findall(r'\d+[.、．]\s*(.+)', title_section.group(1))
        if numbered:
            titles.extend(numbered)
            return titles

    # 格式2: 标题：后面跟内容
    title_matches = re.findall(r'标题[：:]\s*(.+)', content)
    if title_matches:
        titles.extend(title_matches)

    # 格式3: 用户粘贴的多行标题（检测开头几行的短文本）
    lines = content.strip().split('\n')
    if not titles:
        short_lines = []
        for line in lines[:10]:
            line = line.strip()
            if line and len(line) < 50 and not line.startswith('#') and not line.startswith('标签'):
                short_lines.append(line)
            elif short_lines:
                break
        if len(short_lines) >= 2:
            titles = short_lines

    if not titles:
        first_line = extract_title(content)
        if first_line:
            titles = [first_line]

    return titles

def check_forbidden_word(content, word):
    """检查禁词是否出现，返回违规位置列表"""
    exceptions = FORBIDDEN_EXCEPTIONS.get(word, [])
    violations = []
    start = 0
    while True:
        idx = content.find(word, start)
        if idx == -1:
            break
        ctx = content[max(0, idx - 15):idx + len(word) + 15]
        is_exception = any(exc in ctx for exc in exceptions)
        if not is_exception:
            violations.append({"pos": idx, "context": ctx})
        start = idx + 1
    return violations

def auto_insert_fixed_phrases(content):
    """自动插入缺失的不可修改话术，返回修复后的内容"""
    missing_by_cat = {"防敏-水解技术": [], "自护力": [], "基础营养": []}
    for item in FIXED_SELLING_POINTS:
        if item["text"] not in content:
            missing_by_cat[item["category"]].append(item["text"])

    total_missing = sum(len(v) for v in missing_by_cat.values())
    if total_missing == 0:
        return content, 0

    body_match = re.search(r'###\s*正文[^#]*?\n(.*?)(?=###|$)', content, re.DOTALL)
    if not body_match:
        body_match = re.search(r'([\s\S]+)', content)

    if not body_match:
        return content, 0

    body = body_match.group(1)
    modified_body = body
    inserted = 0

    category_anchors = {
        "防敏-水解技术": ["水解", "防敏", "蛋白", "GINI", "致敏"],
        "自护力": ["自护", "HMO", "双菌", "保护力", "菌菌", "肚肚"],
        "基础营养": ["营养", "维生素", "乳糖", "口味"],
    }

    for cat, missing_phrases in missing_by_cat.items():
        if not missing_phrases:
            continue

        anchors = category_anchors.get(cat, [])
        best_pos = -1
        for anchor in anchors:
            pos = modified_body.find(anchor)
            if pos != -1:
                end_pos = modified_body.find("。", pos)
                if end_pos == -1:
                    end_pos = modified_body.find("！", pos)
                if end_pos == -1:
                    end_pos = modified_body.find("\n", pos)
                if end_pos != -1:
                    best_pos = end_pos + 1
                    break

        if best_pos == -1:
            best_pos = len(modified_body)

        for phrase in missing_phrases:
            insert_text = phrase
            if best_pos > 0 and modified_body[best_pos-1] not in "。！\n":
                insert_text = "。" + insert_text
            if not insert_text.endswith(("。", "！")):
                insert_text += "。"

            modified_body = modified_body[:best_pos] + insert_text + modified_body[best_pos:]
            best_pos += len(insert_text)
            inserted += 1

    result = content.replace(body, modified_body)
    return result, inserted

def run_all_checks(content):
    """运行全部审核检查，返回结果字典"""
    results = {}
    title = extract_title(content)
    tags = extract_tags(content)
    word_count = count_chinese(content)

    # 审核1: 卖点顺序
    positions = {}
    for cat, phrases in ORDER_ANCHORS.items():
        min_pos = float('inf')
        for p in phrases:
            idx = content.find(p)
            if idx != -1 and idx < min_pos:
                min_pos = idx
        positions[cat] = min_pos if min_pos != float('inf') else -1
    cats = ["防敏-水解技术", "自护力", "基础营养"]
    order_ok = True
    order_details = []
    for i, cat in enumerate(cats):
        pos = positions[cat]
        found = pos != -1
        order_details.append({"category": cat, "position": pos, "found": found})
        if not found:
            order_ok = False
        elif i > 0 and positions[cats[i - 1]] != -1 and pos < positions[cats[i - 1]]:
            order_ok = False
    results["check1"] = {"status": "pass" if order_ok else "fail", "details": order_details}

    # 审核2: 字数（800-900字）
    results["check2"] = {"status": "pass" if 800 <= word_count <= 900 else "fail", "count": word_count}

    # 审核3: 标题数量（智能检测）
    detected_titles = detect_titles(content)
    title_count = len(detected_titles)
    if title_count >= 3:
        results["check3"] = {"status": "pass", "count": title_count, "titles": detected_titles}
    else:
        results["check3"] = {"status": "fail", "count": title_count, "title": title,
                             "titles": detected_titles,
                             "note": f"当前{title_count}个标题，需提供3个备选"}

    # 审核4: 标签
    missing_tags = [t for t in REQUIRED_TAGS if t not in tags]
    results["check4"] = {
        "status": "pass" if len(tags) >= 10 and not missing_tags else "fail",
        "count": len(tags), "missing": missing_tags, "tags": tags,
    }

    # 审核5: 关键词
    kw_items = []
    all_titles_text = " ".join(detected_titles) if detected_titles else title
    for w in TITLE_KEYWORDS:
        kw_items.append({"scope": "标题", "word": w, "found": w in all_titles_text})
    for w in BODY_KEYWORDS:
        kw_items.append({"scope": "正文", "word": w, "found": w in content})
    for w in COVER_KEYWORDS:
        kw_items.append({"scope": "封面(需人工确认)", "word": w, "found": w in content})
    results["check5"] = {
        "status": "pass" if all(r["found"] for r in kw_items) else "fail",
        "items": kw_items,
    }

    # 审核6: 禁词
    fw_items = []
    for cat, words in FORBIDDEN_WORDS.items():
        for w in words:
            violations = check_forbidden_word(content, w)
            rep = FORBIDDEN_REPLACEMENTS.get(w, "删除")
            fw_items.append({
                "category": cat, "word": w,
                "found": len(violations) > 0,
                "violations": violations,
                "replacement": rep,
            })
    results["check6"] = {
        "status": "fail" if any(r["found"] for r in fw_items) else "pass",
        "items": fw_items,
    }

    # 审核7: 必提需润色卖点
    pp_items = []
    for sp in PARAPHRASE_SELLING_POINTS:
        found = sp["fragment"] in content
        pp_items.append({**sp, "found": found})
    results["check7"] = {
        "status": "pass" if all(r["found"] for r in pp_items) else "fail",
        "items": pp_items,
    }

    # 审核8: 必提不可修改卖点
    fp_items = []
    for sp in FIXED_SELLING_POINTS:
        found = sp["text"] in content
        fp_items.append({**sp, "found": found})
    results["check8"] = {
        "status": "pass" if all(r["found"] for r in fp_items) else "fail",
        "items": fp_items,
    }

    # 审核9: 允许删减的卖点
    op_items = []
    for sp in OPTIONAL_SELLING_POINTS:
        found = sp["fragment"] in content
        op_items.append({**sp, "found": found})
    results["check9"] = {"items": op_items}

    return results

def apply_adopted_changes(original, adopted_map, edit_map, check_results):
    """根据采纳的修改建议生成修改后的文本"""
    modified = original
    changes = []

    if "check6" in check_results:
        for i, item in enumerate(check_results["check6"]["items"]):
            key = f"c6_{i}"
            if adopted_map.get(key) and item["found"]:
                old_word = item["word"]
                new_word = edit_map.get(key, item["replacement"])
                if old_word in modified:
                    exceptions = FORBIDDEN_EXCEPTIONS.get(old_word, [])
                    if exceptions:
                        result = []
                        start = 0
                        while True:
                            idx = modified.find(old_word, start)
                            if idx == -1:
                                result.append(modified[start:])
                                break
                            ctx = modified[max(0, idx - 15):idx + len(old_word) + 15]
                            is_exc = any(exc in ctx for exc in exceptions)
                            if is_exc:
                                result.append(modified[start:idx + len(old_word)])
                            else:
                                result.append(modified[start:idx])
                                result.append(new_word)
                                changes.append({"old": old_word, "new": new_word})
                            start = idx + len(old_word)
                        modified = "".join(result)
                    else:
                        modified = modified.replace(old_word, new_word)
                        changes.append({"old": old_word, "new": new_word})

    if "check4" in check_results:
        missing = check_results["check4"].get("missing", [])
        for i, tag in enumerate(missing):
            key = f"c4_{i}"
            if adopted_map.get(key):
                if tag not in modified:
                    modified = modified.rstrip() + " " + tag
                    changes.append({"old": "", "new": tag})

    return modified, changes

def highlight_diff(text, changes, mode="original"):
    """对文本中的修改部分进行高亮"""
    html = text
    for c in changes:
        if mode == "original" and c["old"]:
            html = html.replace(
                c["old"],
                f'<span style="background:#c8e6c9;padding:1px 4px;border-radius:3px;font-weight:bold;">{c["old"]}</span>'
            )
        elif mode == "modified" and c["new"]:
            html = html.replace(
                c["new"],
                f'<span style="background:#f8bbd0;padding:1px 4px;border-radius:3px;font-weight:bold;">{c["new"]}</span>'
            )
    return html.replace('\n', '<br>')

# ========== 页面配置 ==========
st.set_page_config(page_title="赞意AI审稿系统", page_icon="🤖", layout="wide")

st.markdown("""
<style>
.block-container {padding-top: 1rem !important; padding-bottom: 1rem !important;}
.nav-bar {
    display: flex; gap: 0; margin-bottom: 20px; border-radius: 10px; overflow: hidden;
    border: 2px solid #ddd;
}
.nav-item {
    flex: 1; text-align: center; padding: 12px 10px; font-weight: bold; font-size: 15px;
    cursor: default;
}
.nav-part1 { background: #eef1fa; color: #2c3e6b; border-right: 2px solid #ddd; }
.nav-part2 { background: #f0faf4; color: #2e7d32; border-right: 2px solid #ddd; }
.nav-part3 { background: #fff8e1; color: #f57c00; border-right: 2px solid #ddd; }
.nav-part4 { background: #fce4ec; color: #c2185b; }
[data-testid="stVerticalBlockBorderWrapper"] {
    background-color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part1-marker) {
    background-color: #f5f3ff !important;
    border: 2px solid #c4b5fd !important;
    border-radius: 12px !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part1-marker) button[kind="primary"] {
    background-color: #7c3aed !important; border-color: #7c3aed !important; color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part1-marker) button[kind="primary"]:hover {
    background-color: #6d28d9 !important; border-color: #6d28d9 !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part2-marker) {
    background-color: #edf7f0 !important;
    border: 2px solid #b4dfc6 !important;
    border-radius: 12px !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part2-marker) button[kind="primary"] {
    background-color: #4caf50 !important; border-color: #4caf50 !important; color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part2-marker) button[kind="primary"]:hover {
    background-color: #388e3c !important; border-color: #388e3c !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part3-marker) {
    background-color: #fff8e1 !important;
    border: 2px solid #ffcc80 !important;
    border-radius: 12px !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part3-marker) button[kind="primary"] {
    background-color: #ff9800 !important; border-color: #ff9800 !important; color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part3-marker) button[kind="primary"]:hover {
    background-color: #f57c00 !important; border-color: #f57c00 !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part4-marker) {
    background-color: #fce4ec !important;
    border: 2px solid #f48fb1 !important;
    border-radius: 12px !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part4-marker) button[kind="primary"] {
    background-color: #e91e63 !important; border-color: #e91e63 !important; color: white !important;
}
[data-testid="stVerticalBlockBorderWrapper"]:has(#part4-marker) button[kind="primary"]:hover {
    background-color: #c2185b !important; border-color: #c2185b !important;
}
.check-header-pass {
    background: #e8f5e9; border-left: 5px solid #4caf50; padding: 10px 15px;
    margin: 10px 0 5px 0; border-radius: 0 8px 8px 0; font-weight: bold; font-size: 15px;
}
.check-header-fail {
    background: #fce4ec; border-left: 5px solid #e57373; padding: 10px 15px;
    margin: 10px 0 5px 0; border-radius: 0 8px 8px 0; font-weight: bold; font-size: 15px;
}
.check-header-info {
    background: #fff8e1; border-left: 5px solid #ffc107; padding: 10px 15px;
    margin: 10px 0 5px 0; border-radius: 0 8px 8px 0; font-weight: bold; font-size: 15px;
}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] p {font-size: 0 !important;}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] p::after {content: "将文件拖到此处上传"; font-size: 14px !important;}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] button {font-size: 0 !important; position: relative;}
[data-testid="stFileUploader"] [data-testid="stFileUploaderDropzone"] button::after {content: "选择文件"; font-size: 14px !important; position: absolute;}
[data-testid="stDownloadButton"] > button {
    background-color: #7c3aed !important;
    border-color: #7c3aed !important;
    color: white !important;
}
[data-testid="stDownloadButton"] > button:hover {
    background-color: #6d28d9 !important;
    border-color: #6d28d9 !important;
}
</style>
""", unsafe_allow_html=True)

# ========== 标题 ==========
st.markdown("""
<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 12px; padding: 20px 25px; margin-bottom: 15px;">
    <h2 style="color: white; margin: 0;">🤖 赞意AI · 小红书KOL审稿系统 for 兔子🐰</h2>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="nav-bar">
    <div class="nav-item nav-part1">Part 1 · 八大审核</div>
    <div class="nav-item nav-part2">Part 2 · 人话修改</div>
    <div class="nav-item nav-part3">Part 3 · 复核检查</div>
    <div class="nav-item nav-part4">Part 4 · 终稿完成</div>
</div>
""", unsafe_allow_html=True)

# ========== Session State 初始化 ==========
for key, default in [
    ('kol_content', ''), ('audit_results', None), ('audit_adopted', {}),
    ('audit_edits', {}), ('modified_content', ''), ('diff_changes', []),
    ('renhua_result', ''), ('renhua_adopted', False), ('recheck_content', ''),
    ('recheck_results', None), ('final_content', ''), ('final_ready', False),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ========== 稿件方向选择 ==========
DIRECTION_OPTIONS = [
    "请选择稿件方向...",
    "方向1.【育婴师防敏科普】",
    "方向2.【单品分享】",
    "方向3.【反向经验分享-家族过敏史】",
    "方向4.【反向经验分享-剖腹产】",
    "方向5.【防敏待产包分享-孕晚敏感】",
    "方向6.【防敏待产包分享-剖腹产】",
    "方向7.【养宠家庭】",
    "方向8.【a2VS第一口顶配】",
    "方向9.【能恩全护贵有所值】",
    "方向10.【能恩全护+超启能恩家族测评】",
    "方向11.【防敏竞品测评】",
    "方向12.【跨境能恩全测评】",
]

if 'selected_direction' not in st.session_state:
    st.session_state.selected_direction = DIRECTION_OPTIONS[0]

# ========== 输入区 ==========
dir_col, date_col = st.columns([3, 1])
with dir_col:
    selected_dir = st.selectbox("本稿件符合方向", DIRECTION_OPTIONS, key="direction_select")
    st.session_state.selected_direction = selected_dir
with date_col:
    st.caption(f"当前日期: {TODAY}")

upload_col, paste_col = st.columns(2)
with upload_col:
    kol_file = st.file_uploader("上传KOL稿件 (.docx)", type=["docx"], key="kol_file")
with paste_col:
    kol_text = st.text_area("或粘贴稿件内容", height=120, placeholder="在此粘贴KOL稿件内容...", key="kol_text")

if kol_file:
    kol_file.seek(0)
    st.session_state.kol_content = read_docx(kol_file)
elif kol_text:
    st.session_state.kol_content = kol_text

# ================================================================
# Part 1: 八大审核
# ================================================================
with st.container(border=True):
    st.markdown('<div id="part1-marker"></div>', unsafe_allow_html=True)
    st.markdown("#### Part 1 · 八大审核")
    st.caption("本地Python逐项检查，全部结果以表格展示，发现问题可编辑建议并采纳保存")

    if not st.session_state.kol_content:
        st.info("请先上传或粘贴KOL稿件")
    else:
        if st.button("开始八大审核", key="btn_audit", use_container_width=True, type="primary"):
            st.session_state.audit_results = run_all_checks(st.session_state.kol_content)
            st.session_state.audit_adopted = {}
            st.session_state.audit_edits = {}
            st.session_state.modified_content = ""
            st.session_state.diff_changes = []
            st.rerun()

        if st.session_state.audit_results:
            r = st.session_state.audit_results
            content = st.session_state.kol_content

            pass_count = sum(1 for k in ["check1","check2","check3","check4","check5","check6","check7","check8"]
                           if r.get(k, {}).get("status") == "pass")
            fail_count = 8 - pass_count
            m1, m2, m3 = st.columns(3)
            m1.metric("通过", f"{pass_count}/8")
            m2.metric("需修改", f"{fail_count}")
            m3.metric("稿件字数", f"{r['check2']['count']}")

            # 审核1: 卖点顺序
            s1 = r["check1"]["status"]
            icon1 = "✅" if s1 == "pass" else "❌"
            cls1 = "check-header-pass" if s1 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls1}">{icon1} 审核1：卖点顺序（防敏-水解技术 → 自护力 → 基础营养）</div>', unsafe_allow_html=True)
            rows1 = ""
            for d in r["check1"]["details"]:
                found = "✅ 已出现" if d["found"] else "❌ 未出现"
                pos = f"位置: {d['position']}" if d["found"] else "—"
                bg = "#f0fff4" if d["found"] else "#fff5f5"
                rows1 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;">{d["category"]}</td><td style="border:1px solid #ddd;padding:6px 8px;">{found}</td><td style="border:1px solid #ddd;padding:6px 8px;">{pos}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">卖点类别</th><th style="border:1px solid #ddd;padding:8px;">检查结果</th><th style="border:1px solid #ddd;padding:8px;">位置</th></tr></thead>
            <tbody>{rows1}</tbody></table>''', unsafe_allow_html=True)
            if s1 == "fail":
                st.text_input("修改建议", value="请调整段落顺序：先写防敏-水解技术，再写自护力，最后写基础营养", key="edit_c1", disabled=False)
                st.checkbox("采纳", key="adopt_c1", value=True)

            # 审核2: 字数检查
            s2 = r["check2"]["status"]
            icon2 = "✅" if s2 == "pass" else "❌"
            cls2 = "check-header-pass" if s2 == "pass" else "check-header-fail"
            wc = r["check2"]["count"]
            st.markdown(f'<div class="{cls2}">{icon2} 审核2：字数检查（{wc}字，要求800-900字）</div>', unsafe_allow_html=True)
            if s2 == "fail":
                if wc < 800:
                    st.warning(f"字数不足，还需增加约 {800 - wc} 字")
                else:
                    st.warning(f"字数超标，需精简约 {wc - 900} 字")
                wc_hint = "需扩充内容" if wc < 800 else "需精简内容"
                st.text_input("修改建议", value=f"当前{wc}字，{wc_hint}至800-900字", key="edit_c2")
                st.checkbox("采纳", key="adopt_c2", value=True)

            # 审核3: 标题数量
            s3 = r["check3"]["status"]
            icon3 = "✅" if s3 == "pass" else "❌"
            cls3 = "check-header-pass" if s3 == "pass" else "check-header-fail"
            tc3 = r["check3"]["count"]
            st.markdown(f'<div class="{cls3}">{icon3} 审核3：标题数量（当前{tc3}个，需3个备选标题）</div>', unsafe_allow_html=True)

            detected = r["check3"].get("titles", [])
            if detected:
                rows3_title = ""
                for i, t in enumerate(detected):
                    rows3_title += f'<tr><td style="border:1px solid #ddd;padding:6px 8px;">{i+1}</td><td style="border:1px solid #ddd;padding:6px 8px;">{t[:80]}</td></tr>'
                st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
                <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">#</th><th style="border:1px solid #ddd;padding:8px;">检测到的标题</th></tr></thead>
                <tbody>{rows3_title}</tbody></table>''', unsafe_allow_html=True)
            if s3 == "fail":
                st.caption("建议：人话修改阶段AI将自动生成3个备选标题")

            # 审核4: 话题标签
            s4 = r["check4"]["status"]
            icon4 = "✅" if s4 == "pass" else "❌"
            cls4 = "check-header-pass" if s4 == "pass" else "check-header-fail"
            tc4 = r["check4"]["count"]
            st.markdown(f'<div class="{cls4}">{icon4} 审核4：话题标签（当前{tc4}个，要求10个以上）</div>', unsafe_allow_html=True)

            rows4 = ""
            for tag in REQUIRED_TAGS:
                found = tag in r["check4"]["tags"]
                icon = "✅" if found else "❌ 缺失"
                bg = "#f0fff4" if found else "#fff5f5"
                rows4 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;">{tag}</td><td style="border:1px solid #ddd;padding:6px 8px;">{icon}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">必含标签</th><th style="border:1px solid #ddd;padding:8px;">结果</th></tr></thead>
            <tbody>{rows4}</tbody></table>''', unsafe_allow_html=True)

            if r["check4"]["missing"]:
                for mi, mtag in enumerate(r["check4"]["missing"]):
                    c4_col1, c4_col2 = st.columns([3, 1])
                    with c4_col1:
                        st.text_input(f"补充标签", value=mtag, key=f"edit_c4_{mi}")
                    with c4_col2:
                        st.checkbox("采纳", key=f"adopt_c4_{mi}", value=True)

            # 审核5: 关键词
            s5 = r["check5"]["status"]
            icon5 = "✅" if s5 == "pass" else "❌"
            cls5 = "check-header-pass" if s5 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls5}">{icon5} 审核5：必须出现关键词</div>', unsafe_allow_html=True)

            rows5 = ""
            for item in r["check5"]["items"]:
                found = "✅ 已包含" if item["found"] else "❌ 缺失"
                bg = "#f0fff4" if item["found"] else "#fff5f5"
                rows5 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;">{item["scope"]}</td><td style="border:1px solid #ddd;padding:6px 8px;font-weight:bold;">{item["word"]}</td><td style="border:1px solid #ddd;padding:6px 8px;">{found}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">检查范围</th><th style="border:1px solid #ddd;padding:8px;">关键词</th><th style="border:1px solid #ddd;padding:8px;">结果</th></tr></thead>
            <tbody>{rows5}</tbody></table>''', unsafe_allow_html=True)

            missing_kw = [item for item in r["check5"]["items"] if not item["found"]]
            for ki, kw_item in enumerate(missing_kw):
                c5_col1, c5_col2 = st.columns([3, 1])
                with c5_col1:
                    st.text_input(f"修改建议", value=f"请在{kw_item['scope']}中加入「{kw_item['word']}」", key=f"edit_c5_{ki}")
                with c5_col2:
                    st.checkbox("采纳", key=f"adopt_c5_{ki}", value=True)

            # 审核6: 禁词/禁用表达
            s6 = r["check6"]["status"]
            icon6 = "✅" if s6 == "pass" else "❌"
            cls6 = "check-header-pass" if s6 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls6}">{icon6} 审核6：禁词/禁用表达检查</div>', unsafe_allow_html=True)

            rows6 = ""
            for item in r["check6"]["items"]:
                if item["found"]:
                    icon = "❌ 出现了"
                    bg = "#fff5f5"
                    ctx_list = item.get("violations", [])
                    ctx_str = "、".join([f'"{v["context"].strip()}"' for v in ctx_list[:2]])
                else:
                    icon = "✅ 未出现"
                    bg = "#f0fff4"
                    ctx_str = "—"
                rows6 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;">{item["category"]}</td><td style="border:1px solid #ddd;padding:6px 8px;font-weight:bold;">{item["word"]}</td><td style="border:1px solid #ddd;padding:6px 8px;">{icon}</td><td style="border:1px solid #ddd;padding:6px 8px;font-size:12px;">{ctx_str}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">类型</th><th style="border:1px solid #ddd;padding:8px;">禁词</th><th style="border:1px solid #ddd;padding:8px;">结果</th><th style="border:1px solid #ddd;padding:8px;">上下文</th></tr></thead>
            <tbody>{rows6}</tbody></table>''', unsafe_allow_html=True)

            found_forbidden = [item for item in r["check6"]["items"] if item["found"]]
            for fi, fw in enumerate(found_forbidden):
                c6_col1, c6_col2 = st.columns([3, 1])
                with c6_col1:
                    default_rep = fw["replacement"]
                    edited = st.text_input(f"「{fw['word']}」替换为", value=default_rep, key=f"edit_c6_{fi}")
                    st.session_state.audit_edits[f"c6_{r['check6']['items'].index(fw)}"] = edited
                with c6_col2:
                    st.checkbox("采纳", key=f"adopt_c6_{fi}", value=True)

            # 审核7: 必提需润色卖点
            s7 = r["check7"]["status"]
            icon7 = "✅" if s7 == "pass" else "❌"
            cls7 = "check-header-pass" if s7 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls7}">{icon7} 审核7：必提需润色卖点（4大方向 · 10小方向）</div>', unsafe_allow_html=True)

            current_cat7 = ""
            pi_counter = 0
            for item in r["check7"]["items"]:
                if item["category"] != current_cat7:
                    current_cat7 = item["category"]
                    st.markdown(f'<div style="background:#e8eaf6;padding:6px 12px;margin-top:10px;border-radius:5px;font-weight:bold;color:#3949ab;">📂 大方向：{current_cat7}</div>', unsafe_allow_html=True)

                found = item["found"]
                icon = "✅" if found else "❌"
                bg = "#f0fff4" if found else "#fff5f5"
                border_color = "#4caf50" if found else "#ef5350"

                st.markdown(f'''<div style="background:{bg};border-left:4px solid {border_color};padding:10px 15px;margin:6px 0;border-radius:0 8px 8px 0;">
                <div style="font-size:13px;"><b>小方向{item["idx"]}</b> {icon}</div>
                <div style="font-size:13px;color:#333;margin-top:4px;line-height:1.6;">{item["text"]}</div>
                </div>''', unsafe_allow_html=True)

                if not found:
                    c7_col1, c7_col2 = st.columns([4, 1])
                    with c7_col1:
                        st.text_input("修改建议", value=f"需润色加入：{item['text']}", key=f"edit_c7_{pi_counter}", label_visibility="collapsed")
                    with c7_col2:
                        st.checkbox("采纳", key=f"adopt_c7_{pi_counter}", value=True)
                    pi_counter += 1

            # 审核8: 必提不可修改卖点
            s8 = r["check8"]["status"]
            icon8 = "✅" if s8 == "pass" else "❌"
            cls8 = "check-header-pass" if s8 == "pass" else "check-header-fail"
            st.markdown(f'<div class="{cls8}">{icon8} 审核8：必提不可修改卖点（3大切角 · 10小切角，必须字字不差）</div>', unsafe_allow_html=True)

            current_cat8 = ""
            fpi_counter = 0
            for item in r["check8"]["items"]:
                if item["category"] != current_cat8:
                    current_cat8 = item["category"]
                    st.markdown(f'<div style="background:#fce4ec;padding:6px 12px;margin-top:10px;border-radius:5px;font-weight:bold;color:#c2185b;">📂 大切角：{current_cat8}</div>', unsafe_allow_html=True)

                found = item["found"]
                icon = "✅" if found else "❌"
                bg = "#f0fff4" if found else "#fff5f5"
                border_color = "#4caf50" if found else "#ef5350"

                if found:
                    st.markdown(f'''<div style="background:{bg};border-left:4px solid {border_color};padding:10px 15px;margin:6px 0;border-radius:0 8px 8px 0;">
                    <div style="font-size:13px;"><b>小切角{item["idx"]}</b> {icon} <span style="color:#4caf50;font-size:11px;">（已包含）</span></div>
                    <div style="font-size:13px;color:#333;margin-top:4px;line-height:1.6;font-weight:500;">{item["text"]}</div>
                    </div>''', unsafe_allow_html=True)
                else:
                    st.markdown(f'''<div style="background:{bg};border-left:4px solid {border_color};padding:10px 15px;margin:6px 0;border-radius:0 8px 8px 0;">
                    <div style="font-size:13px;"><b>小切角{item["idx"]}</b> {icon} <span style="color:#c62828;font-size:11px;font-weight:bold;">没有提到</span></div>
                    <div style="font-size:13px;color:#c62828;margin-top:4px;line-height:1.6;font-weight:600;">建议增加：<span style="color:#333;">{item["text"]}</span></div>
                    </div>''', unsafe_allow_html=True)

                if not found:
                    c8_col1, c8_col2 = st.columns([4, 1])
                    with c8_col1:
                        st.text_input("修改建议", value=f"必须原封不动加入：{item['text']}", key=f"edit_c8_{fpi_counter}", label_visibility="collapsed")
                    with c8_col2:
                        st.checkbox("采纳", key=f"adopt_c8_{fpi_counter}", value=True)
                    fpi_counter += 1

            # 审核9: 允许删减的卖点
            st.markdown(f'<div class="check-header-info">ℹ️ 审核9：允许删减的卖点（仅供参考）</div>', unsafe_allow_html=True)
            rows9 = ""
            for item in r["check9"]["items"]:
                found = "✅ 已出现" if item["found"] else "— 未出现（可删减）"
                bg = "#f0fff4" if item["found"] else "#fffde7"
                rows9 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:6px 8px;"><b>{item["category"]}</b></td><td style="border:1px solid #ddd;padding:6px 8px;font-size:12px;">{item["text"]}</td><td style="border:1px solid #ddd;padding:6px 8px;">{found}</td></tr>'
            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:4px 0 12px 0;">
            <thead><tr style="background:#f0f2f6;"><th style="border:1px solid #ddd;padding:8px;">类别</th><th style="border:1px solid #ddd;padding:8px;">卖点内容</th><th style="border:1px solid #ddd;padding:8px;">状态</th></tr></thead>
            <tbody>{rows9}</tbody></table>''', unsafe_allow_html=True)

            with st.expander("📖 标准卖点示例（参考）", expanded=False):
                st.markdown(SELLING_POINT_EXAMPLE)

            st.markdown("---")
            if st.button("保存所有采纳修改 → 生成对比预览", key="btn_save_audit", use_container_width=True, type="primary"):
                adopted = {}
                edits = {}

                found_fw = [item for item in r["check6"]["items"] if item["found"]]
                for fi, fw in enumerate(found_fw):
                    real_idx = r["check6"]["items"].index(fw)
                    adopted[f"c6_{real_idx}"] = st.session_state.get(f"adopt_c6_{fi}", False)
                    edits[f"c6_{real_idx}"] = st.session_state.get(f"edit_c6_{fi}", fw["replacement"])

                missing_tags = r["check4"].get("missing", [])
                for mi, _ in enumerate(missing_tags):
                    adopted[f"c4_{mi}"] = st.session_state.get(f"adopt_c4_{mi}", False)

                st.session_state.audit_adopted = adopted
                st.session_state.audit_edits = edits

                modified, changes = apply_adopted_changes(
                    st.session_state.kol_content, adopted, edits, r
                )
                st.session_state.modified_content = modified
                st.session_state.diff_changes = changes
                st.rerun()

            if st.session_state.modified_content:
                st.markdown("---")
                st.markdown("### 对比预览（原文 vs 修改后）")
                st.caption("🟢 绿色 = 原文中被修改的部分 | 🩷 粉色 = 修改后的内容")

                cmp_left, cmp_right = st.columns(2)
                with cmp_left:
                    st.markdown("**原文（绿色标注修改处）**")
                    orig_html = highlight_diff(st.session_state.kol_content, st.session_state.diff_changes, "original")
                    st.markdown(f'<div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:15px;font-size:14px;line-height:2.0;">{orig_html}</div>', unsafe_allow_html=True)

                with cmp_right:
                    st.markdown("**修改后（粉色标注修改处）**")
                    mod_html = highlight_diff(st.session_state.modified_content, st.session_state.diff_changes, "modified")
                    st.markdown(f'<div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:15px;font-size:14px;line-height:2.0;">{mod_html}</div>', unsafe_allow_html=True)

                with st.expander("需要微调？点击编辑修改后内容", expanded=False):
                    edited_mod = st.text_area("修改后内容（可编辑）", st.session_state.modified_content, height=300, key="edit_modified")
                    if edited_mod != st.session_state.modified_content:
                        st.session_state.modified_content = edited_mod

                adopt_col, dl_col = st.columns(2)
                with adopt_col:
                    if st.button("采用修改后稿件（进入人话修改）", key="btn_adopt_audit", use_container_width=True, type="primary"):
                        st.session_state.kol_content = st.session_state.modified_content
                        st.success("已采用！可进入下方人话修改")

                with dl_col:
                    from docx.shared import Pt
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = 'PingFang SC'
                    style.font.size = Pt(11)
                    output_name = f"采纳后稿件_{TODAY}"
                    doc.add_heading("采纳后稿件", 0)
                    for line in st.session_state.modified_content.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button("📥 下载采纳后稿件", buf, f"{output_name}.docx",
                                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                     key="dl_audit")

# ================================================================
# Part 2: 人话修改
# ================================================================
with st.container(border=True):
    st.markdown('<div id="part2-marker"></div>', unsafe_allow_html=True)
    st.markdown("#### Part 2 · 人话修改（六步审计法）")
    st.caption("AI按照六步审计法对稿件进行人话修改：卖点逻辑→结构完整性→口吻人设→关键词禁词→话术回填→内容结构占比")

    if not st.session_state.kol_content:
        st.info("请先上传稿件并完成八大审核")
    else:
        st.markdown(f'<div style="background:#fff;border-left:3px solid #4caf50;padding:8px 12px;font-size:13px;margin-bottom:10px;">当前稿件：{count_chinese(st.session_state.kol_content)} 字</div>', unsafe_allow_html=True)

        if st.button("开始人话修改（自动循环至八大审核全通过）", key="btn_renhua", use_container_width=True, type="primary"):
            max_retries = 5
            retry_count = 0
            current_content = st.session_state.kol_content
            final_result = None
            all_passed = False

            progress_bar = st.progress(0)
            status_text = st.empty()
            detail_text = st.empty()

            while retry_count < max_retries and not all_passed:
                retry_count += 1
                status_text.markdown(f"🔄 **第 {retry_count} 次生成中...**（最多尝试{max_retries}次）")
                progress_bar.progress(retry_count / max_retries * 0.8)

                if retry_count == 1:
                    prompt = RENHUA_PROMPT.replace("{content}", current_content)
                else:
                    fix_hints = []
                    r_loop = run_all_checks(final_result)
                    if r_loop.get("check1", {}).get("status") != "pass":
                        fix_hints.append("- 调整卖点顺序：必须按 防敏-水解技术→自护力→基础营养 顺序")
                    if r_loop.get("check2", {}).get("status") != "pass":
                        wc_loop = r_loop['check2']['count']
                        hint_loop = "字数不足，需扩充" if wc_loop < 800 else "字数超标，需精简"
                        fix_hints.append(f"- {hint_loop}：当前{wc_loop}字，必须在800-900字之间")
                    if r_loop.get("check3", {}).get("status") != "pass":
                        fix_hints.append("- 必须提供3个备选标题（格式：### 标题备选（3个）然后 1. 2. 3.）")
                    if r_loop.get("check4", {}).get("status") != "pass":
                        missing_loop = r_loop['check4'].get('missing', [])
                        fix_hints.append(f"- 补充标签：{', '.join(missing_loop)}")
                    if r_loop.get("check5", {}).get("status") != "pass":
                        fix_hints.append("- 标题必含【适度水解、防敏、科普】，正文必含【适度水解、防敏、能恩全护】")
                    if r_loop.get("check6", {}).get("status") != "pass":
                        found_loop = [x['word'] for x in r_loop['check6']['items'] if x['found']]
                        rep_loop = {x['word']: x['replacement'] for x in r_loop['check6']['items'] if x['found']}
                        fix_hints.append(f"- 替换禁词：" + "、".join([f"{w}→{rep_loop[w]}" for w in found_loop]))
                    if r_loop.get("check7", {}).get("status") != "pass":
                        missing7_loop = [x['fragment'] for x in r_loop['check7']['items'] if not x['found']]
                        fix_hints.append(f"- 补充润色卖点关键词：{', '.join(missing7_loop[:5])}")
                    if r_loop.get("check8", {}).get("status") != "pass":
                        missing8_loop = [x['text'] for x in r_loop['check8']['items'] if not x['found']]
                        fix_hints.append(f"- 必须原封不动加入以下话术：\n  " + "\n  ".join(missing8_loop))

                    detail_text.markdown("**当前未通过项：**\n" + "\n".join(fix_hints))

                    fix_text = "\n".join(fix_hints)
                    prompt = f"""请修正以下稿件，解决检测到的问题：

【需要修正的问题】
{fix_text}

【当前稿件】
{final_result}

【修正要求】
1. ⚠️ 正文必须在800-900字之间（最重要！）
2. 必须提供3个备选标题（包含：适度水解、防敏、科普）
3. 必须包含10个以上标签，包括：#能恩全护 #适度水解 #适度水解奶粉推荐 #第一口奶粉
4. 替换所有禁词（敏宝→敏感体质宝宝、过敏→敏敏、新生儿→初生宝宝、预防→防敏、生长→成长、发育→噌噌长、免疫→保护力、疾病→不适）
5. 绝对禁止出现：敏宝、奶瓶、奶嘴
6. 必须包含全部10句不可修改话术（字字不差）
7. 保持小红书活人感爆文风格

请直接输出修正后的完整稿件：
### 标题备选（3个）
### 正文（800-900字，必须写够！）
### 话题标签（10个以上）"""

                result = call_llm_api(prompt)
                if result and not result.startswith("Error"):
                    final_result = result
                    final_result, inserted_count = auto_insert_fixed_phrases(final_result)
                    if inserted_count > 0:
                        status_text.markdown(f"📝 第{retry_count}次 - 自动补充了 {inserted_count} 条缺失话术")
                    check_result = run_all_checks(final_result)
                    pass_count = sum(1 for k in ["check1","check2","check3","check4","check5","check6","check7","check8"]
                                   if check_result.get(k, {}).get("status") == "pass")
                    status_text.markdown(f"🔍 第 {retry_count} 次检查：通过 {pass_count}/8 项")
                    if pass_count == 8:
                        all_passed = True
                else:
                    status_text.error(f"AI调用失败: {result}")
                    break

            progress_bar.progress(1.0)
            if all_passed:
                status_text.success(f"✅ 八大审核全部通过！（共尝试 {retry_count} 次）")
                detail_text.empty()
                st.session_state.renhua_result = final_result
                st.session_state.recheck_results = run_all_checks(final_result)
                st.rerun()
            elif final_result:
                status_text.warning(f"⚠️ 已达最大尝试次数({max_retries}次)，当前结果可能仍有未通过项，可手动编辑修正")
                st.session_state.renhua_result = final_result
                st.session_state.recheck_results = run_all_checks(final_result)
                st.rerun()

        if st.session_state.renhua_result:
            st.markdown("---")

            with st.expander("📄 审核后稿件（修改前）", expanded=False):
                orig_html = st.session_state.kol_content.replace('\n', '<br>')
                st.markdown(f'<div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:15px;font-size:13px;line-height:1.8;">{orig_html}</div>', unsafe_allow_html=True)

            st.markdown("### 🔍 复核检查（人话修改后自动验证）")
            result_text = st.session_state.renhua_result

            title_matches = re.findall(r'###\s*标题备选.*?(?=###|$)', result_text, re.DOTALL)
            title_section = title_matches[0] if title_matches else ""
            title_count = len(re.findall(r'\d+\.\s*.+', title_section))

            body_matches = re.findall(r'###\s*正文.*?(?=###|$)', result_text, re.DOTALL)
            body_section = body_matches[0] if body_matches else result_text
            body_word_count = count_chinese(body_section)

            tags_in_result = extract_tags(result_text)

            human_markers = ["我", "你", "咱", "说实话", "不瞒你说", "一开始", "其实", "真的", "姐妹", "绝了", "救命", "后悔"]
            human_found = sum(1 for m in human_markers if m in result_text)
            emoji_markers = ["💡", "✨", "🔥", "❗", "👶", "🍼", "💪", "❤️", "🙋", "😊"]
            emoji_found = sum(1 for e in emoji_markers if e in result_text)
            exclamation_count = result_text.count("！") + result_text.count("!")

            check_items = [
                ("审核2 - 字数", f"{body_word_count}字（要求800-900）", "pass" if 800 <= body_word_count <= 900 else "fail"),
                ("审核3 - 标题数量", f"{title_count}个备选标题", "pass" if title_count >= 3 else "fail"),
                ("审核4 - 标签数量", f"{len(tags_in_result)}个标签", "pass" if len(tags_in_result) >= 10 else "fail"),
                ("活人感关键词", f"包含{human_found}/{len(human_markers)}个口语化表达", "pass" if human_found >= 5 else "warn"),
                ("Emoji使用", f"包含{emoji_found}个emoji", "pass" if emoji_found >= 3 else "warn"),
                ("爆文语气", f"{exclamation_count}个感叹号", "pass" if exclamation_count >= 5 else "warn"),
            ]

            rows_recheck = ""
            for name, detail, status in check_items:
                if status == "pass":
                    icon = "✅"
                    bg = "#f0fff4"
                elif status == "fail":
                    icon = "❌"
                    bg = "#fff5f5"
                else:
                    icon = "⚠️"
                    bg = "#fffde7"
                rows_recheck += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:8px;">{name}</td><td style="border:1px solid #ddd;padding:8px;">{detail}</td><td style="border:1px solid #ddd;padding:8px;font-weight:bold;">{icon}</td></tr>'

            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:8px 0 16px 0;">
            <thead><tr style="background:#e3f2fd;"><th style="border:1px solid #ddd;padding:8px;">检查项</th><th style="border:1px solid #ddd;padding:8px;">详情</th><th style="border:1px solid #ddd;padding:8px;">结果</th></tr></thead>
            <tbody>{rows_recheck}</tbody></table>''', unsafe_allow_html=True)

            with st.expander("📝 小红书爆文笔记攻略 · 活人感检查", expanded=True):
                st.markdown("**🗣️ 口语化表达**")
                markers_detail = []
                for m in human_markers:
                    if m in result_text:
                        markers_detail.append(f'<span style="background:#c8e6c9;padding:2px 6px;border-radius:3px;margin:2px;">✅ {m}</span>')
                    else:
                        markers_detail.append(f'<span style="background:#ffcdd2;padding:2px 6px;border-radius:3px;margin:2px;">❌ {m}</span>')
                st.markdown(f'<div style="line-height:2.2;">{"".join(markers_detail)}</div>', unsafe_allow_html=True)

                st.markdown("**😊 Emoji使用**")
                emoji_detail = []
                for e in emoji_markers:
                    if e in result_text:
                        emoji_detail.append(f'<span style="background:#c8e6c9;padding:2px 6px;border-radius:3px;margin:2px;">✅ {e}</span>')
                    else:
                        emoji_detail.append(f'<span style="background:#ffcdd2;padding:2px 6px;border-radius:3px;margin:2px;">❌ {e}</span>')
                st.markdown(f'<div style="line-height:2.2;">{"".join(emoji_detail)}</div>', unsafe_allow_html=True)

                st.markdown(f"**🔥 爆文语气**：共{exclamation_count}个感叹号（建议≥5个）")
                st.caption("小红书爆文特征：多用感叹号、emoji、口语化表达，像闺蜜聊天一样自然")

            st.markdown("---")
            st.markdown("### 人话修改结果")
            st.markdown(st.session_state.renhua_result)

            with st.expander("需要微调？点击编辑", expanded=False):
                edited_renhua = st.text_area("人话修改内容（可编辑）", st.session_state.renhua_result, height=400, key="edit_renhua")
                if edited_renhua != st.session_state.renhua_result:
                    st.session_state.renhua_result = edited_renhua

            if st.button("采用人话修改结果 → 进入Part 3复核", key="btn_adopt_renhua", use_container_width=True, type="primary"):
                st.session_state.renhua_adopted = True
                st.session_state.recheck_content = st.session_state.renhua_result
                st.session_state.recheck_results = None
                st.session_state.final_ready = False
                st.success("已采用！请在下方Part 3进行复核检查")
                st.rerun()

# ================================================================
# Part 3: 复核检查
# ================================================================
with st.container(border=True):
    st.markdown('<div id="part3-marker"></div>', unsafe_allow_html=True)
    st.markdown("#### Part 3 · 复核检查（再次八大审核）")
    st.caption("对人话修改后的稿件进行八大审核，确保合规后可编辑微调")

    if not st.session_state.renhua_adopted or not st.session_state.recheck_content:
        st.info("请先完成Part 2人话修改并采用结果")
    else:
        recheck_wc = count_chinese(st.session_state.recheck_content)
        st.markdown(f'<div style="background:#fff;border-left:3px solid #ff9800;padding:8px 12px;font-size:13px;margin-bottom:10px;">待复核稿件：{recheck_wc} 字</div>', unsafe_allow_html=True)

        if st.button("开始复核（八大审核）", key="btn_recheck", use_container_width=True, type="primary"):
            st.session_state.recheck_results = run_all_checks(st.session_state.recheck_content)
            st.rerun()

        if st.session_state.recheck_results:
            r3 = st.session_state.recheck_results

            pass_count3 = sum(1 for k in ["check1","check2","check3","check4","check5","check6","check7","check8"]
                           if r3.get(k, {}).get("status") == "pass")
            fail_count3 = 8 - pass_count3

            st.markdown("### 复核结果概览")
            m3_1, m3_2, m3_3 = st.columns(3)
            m3_1.metric("通过", f"{pass_count3}/8", delta="良好" if pass_count3 >= 6 else "需修改")
            m3_2.metric("需修改", f"{fail_count3}")
            m3_3.metric("字数", f"{r3['check2']['count']}字", delta="800-900" if 800 <= r3['check2']['count'] <= 900 else "需调整")

            st.markdown("### 八大审核结果")
            check_names = [
                ("check1", "审核1-卖点顺序"),
                ("check2", "审核2-字数检查"),
                ("check3", "审核3-标题数量"),
                ("check4", "审核4-话题标签"),
                ("check5", "审核5-关键词"),
                ("check6", "审核6-禁词检查"),
                ("check7", "审核7-润色卖点"),
                ("check8", "审核8-不可修改卖点"),
            ]
            rows3 = ""
            for key, name in check_names:
                status = r3.get(key, {}).get("status", "fail")
                icon = "✅" if status == "pass" else "❌"
                bg = "#f0fff4" if status == "pass" else "#fff5f5"
                if key == "check2":
                    detail = f"{r3['check2']['count']}字（要求800-900）"
                elif key == "check3":
                    detail = f"{r3['check3']['count']}个标题"
                elif key == "check4":
                    detail = f"{r3['check4']['count']}个标签，缺失{len(r3['check4']['missing'])}个"
                elif key == "check7":
                    missing7 = len([x for x in r3['check7']['items'] if not x['found']])
                    detail = f"缺失{missing7}/10个润色卖点"
                elif key == "check8":
                    missing8 = len([x for x in r3['check8']['items'] if not x['found']])
                    detail = f"缺失{missing8}/10个不可修改卖点"
                else:
                    detail = "—"
                rows3 += f'<tr style="background:{bg};"><td style="border:1px solid #ddd;padding:8px;">{name}</td><td style="border:1px solid #ddd;padding:8px;">{detail}</td><td style="border:1px solid #ddd;padding:8px;font-weight:bold;">{icon}</td></tr>'

            st.markdown(f'''<table style="width:100%;border-collapse:collapse;font-size:13px;margin:8px 0;">
            <thead><tr style="background:#fff3e0;"><th style="border:1px solid #ddd;padding:8px;">检查项</th><th style="border:1px solid #ddd;padding:8px;">详情</th><th style="border:1px solid #ddd;padding:8px;">结果</th></tr></thead>
            <tbody>{rows3}</tbody></table>''', unsafe_allow_html=True)

            if fail_count3 > 0:
                st.markdown("---")
                st.warning(f"⚠️ 检测到 {fail_count3} 项未通过，需要重新生成人话版本")

                if st.button("🔄 重新生成人话版本（AI自动修正）", key="btn_regenerate", use_container_width=True, type="primary"):
                    with st.spinner("AI重新生成中，自动修正未通过项..."):
                        fix_hints = []
                        if r3.get("check1", {}).get("status") != "pass":
                            fix_hints.append("- 调整卖点顺序：必须按 防敏-水解技术→自护力→基础营养 顺序")
                        if r3.get("check2", {}).get("status") != "pass":
                            wc3 = r3['check2']['count']
                            hint3 = "字数不足，需扩充" if wc3 < 800 else "字数超标，需精简"
                            fix_hints.append(f"- {hint3}：当前{wc3}字，必须在800-900字之间")
                        if r3.get("check3", {}).get("status") != "pass":
                            fix_hints.append("- 补充标题：必须提供3个备选标题（格式：1. 2. 3.）")
                        if r3.get("check4", {}).get("status") != "pass":
                            missing_tags = r3['check4'].get('missing', [])
                            fix_hints.append(f"- 补充标签：缺失 {', '.join(missing_tags[:5])}")
                        if r3.get("check6", {}).get("status") != "pass":
                            found_fw3 = [x['word'] for x in r3['check6']['items'] if x['found']]
                            rep3 = {x['word']: x['replacement'] for x in r3['check6']['items'] if x['found']}
                            fix_hints.append(f"- 替换禁词：" + "、".join([f"{w}→{rep3[w]}" for w in found_fw3]))
                        if r3.get("check7", {}).get("status") != "pass":
                            fix_hints.append("- 补充润色卖点：确保10个小方向都有体现")
                        if r3.get("check8", {}).get("status") != "pass":
                            missing_fixed = [x['text'] for x in r3['check8']['items'] if not x['found']]
                            fix_hints.append(f"- 补充不可修改卖点（必须字字不差）：\n  " + "\n  ".join(missing_fixed))

                        fix_prompt = "\n".join(fix_hints)

                        regen_prompt = f"""你是小红书爆文写手。请修正以下稿件，解决检测到的问题。

【需要修正的问题】
{fix_prompt}

【原稿件】
{st.session_state.recheck_content}

【修正要求 - 按重要性排序】
1. ⚠️ 正文必须在800-900字之间（最重要！写够字数！）
2. 必须有小红书活人感（用"姐妹们""真的绝了""说实话"等口语化表达）
3. 必须提供3个备选标题（包含：适度水解、防敏、科普）
4. 必须包含10个以上标签（必含 #能恩全护 #适度水解）
5. 替换禁词（敏宝→敏感体质宝宝、过敏→敏敏、新生儿→初生宝宝、预防→防敏、免疫→保护力）
6. 必须包含全部10句不可修改话术（字字不差）

请直接输出修正后的完整稿件：
### 标题备选（3个）
### 正文（800-900字，必须写够！）
### 话题标签（10个以上）"""

                        result = call_llm_api(regen_prompt)
                        if result and not result.startswith("Error"):
                            result, inserted_count = auto_insert_fixed_phrases(result)
                            if inserted_count > 0:
                                st.info(f"📝 自动补充了 {inserted_count} 条缺失话术")
                            st.session_state.recheck_content = result
                            st.session_state.edit_recheck_content = result
                            st.session_state.recheck_results = run_all_checks(result)
                            st.rerun()
                        else:
                            st.error(f"AI调用失败: {result}")

                st.markdown("---")
                st.markdown("### 或手动编辑修正")
            else:
                st.success("🎉 恭喜！八大审核全部通过！")
                st.markdown("---")
                st.markdown("### 最终稿件预览")

            edited_recheck = st.text_area(
                "编辑正文内容",
                st.session_state.recheck_content,
                height=400,
                key="edit_recheck_content"
            )
            if edited_recheck != st.session_state.recheck_content:
                st.session_state.recheck_content = edited_recheck

            current_wc = count_chinese(edited_recheck)
            wc_color = "#4caf50" if 800 <= current_wc <= 900 else "#f44336"
            st.markdown(f'<div style="text-align:right;color:{wc_color};font-weight:bold;">当前字数：{current_wc}/900</div>', unsafe_allow_html=True)

            if fail_count3 > 0:
                if st.button("🔍 重新检查（手动修改后）", key="btn_manual_recheck", use_container_width=True):
                    st.session_state.recheck_results = run_all_checks(st.session_state.recheck_content)
                    st.rerun()

            if fail_count3 == 0:
                if st.button("✅ 确认复核完成 → 进入Part 4终稿", key="btn_confirm_recheck", use_container_width=True, type="primary"):
                    st.session_state.final_content = st.session_state.recheck_content
                    st.session_state.final_ready = True
                    st.success("复核完成！请在Part 4预览并下载终稿")
                    st.rerun()
            else:
                st.info("💡 请先修正所有未通过项，八大审核全部通过后才能进入Part 4")

# ================================================================
# Part 4: 终稿完成
# ================================================================
with st.container(border=True):
    st.markdown('<div id="part4-marker"></div>', unsafe_allow_html=True)
    st.markdown("#### Part 4 · 终稿完成")
    st.caption("预览终稿并下载")

    if not st.session_state.final_ready or not st.session_state.final_content:
        st.info("请先完成Part 3复核检查")
    else:
        final_wc = count_chinese(st.session_state.final_content)
        final_tags = extract_tags(st.session_state.final_content)
        dir_name = st.session_state.selected_direction if st.session_state.selected_direction != DIRECTION_OPTIONS[0] else "未指定方向"

        st.markdown(f'''
        <div style="background:#fff;border:2px solid #f48fb1;border-radius:10px;padding:15px;margin-bottom:15px;">
            <div style="font-size:16px;font-weight:bold;color:#c2185b;margin-bottom:10px;">📋 终稿信息</div>
            <div style="display:flex;gap:20px;flex-wrap:wrap;">
                <div>📝 字数：<b>{final_wc}</b></div>
                <div>🏷️ 标签：<b>{len(final_tags)}个</b></div>
                <div>📂 方向：<b>{dir_name}</b></div>
                <div>📅 日期：<b>{TODAY}</b></div>
            </div>
        </div>
        ''', unsafe_allow_html=True)

        st.markdown("### 终稿预览")
        final_html = st.session_state.final_content.replace('\n', '<br>')
        st.markdown(f'''<div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:20px;font-size:14px;line-height:2.0;max-height:500px;overflow-y:auto;">
        {final_html}
        </div>''', unsafe_allow_html=True)

        st.markdown("---")
        dl_col1, dl_col2 = st.columns(2)

        with dl_col1:
            from docx.shared import Pt, RGBColor
            doc_final = Document()
            style_final = doc_final.styles['Normal']
            style_final.font.name = 'PingFang SC'
            style_final.font.size = Pt(11)
            output_name_final = f"KOL_{TODAY}_终稿"
            doc_final.add_heading(output_name_final, 0)
            doc_final.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc_final.add_paragraph(f"稿件方向: {dir_name}")
            doc_final.add_paragraph(f"字数: {final_wc}")
            doc_final.add_paragraph("─" * 50)
            for line in st.session_state.final_content.split('\n'):
                if line.strip():
                    doc_final.add_paragraph(line.strip())
            buf_final = io.BytesIO()
            doc_final.save(buf_final)
            buf_final.seek(0)
            st.download_button(
                "📥 下载终稿 (.docx)",
                buf_final,
                f"{output_name_final}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_final",
                use_container_width=True
            )

        with dl_col2:
            st.download_button(
                "📄 下载纯文本 (.txt)",
                st.session_state.final_content,
                f"KOL_{TODAY}_终稿.txt",
                "text/plain",
                key="dl_final_txt",
                use_container_width=True
            )

        st.success("🎉 恭喜！终稿已完成，可下载使用")

# ========== Footer ==========
st.markdown("---")
dir_label = st.session_state.selected_direction if st.session_state.selected_direction != DIRECTION_OPTIONS[0] else "能恩全护"
st.caption(f"🤖 赞意AI审稿系统 v4.1 · {dir_label}")
